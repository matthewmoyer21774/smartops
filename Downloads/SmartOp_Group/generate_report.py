"""
Generate a Word document explaining the inventory management methodology.
Run: python generate_report.py
Output: Methodology_Report.docx
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# --- Style setup ---
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.color.rgb = RGBColor(0, 90, 156)
    h.font.name = "Calibri"

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph("")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Data-Driven Inventory Management")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 90, 156)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Methodology Report")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph("")

course = doc.add_paragraph()
course.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = course.add_run("SmartOp - Final Assignment\nKU Leuven")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (placeholder)
# ============================================================
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Problem Overview",
    "2. Data Exploration",
    "3. Demand Forecasting Model",
    "4. Inventory Simulation Engine",
    "5. Ordering Policy",
    "6. Backtest Results",
    "7. How to Run the Live Game",
    "8. Key Design Decisions",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. PROBLEM OVERVIEW
# ============================================================
doc.add_heading("1. Problem Overview", level=1)

doc.add_paragraph(
    "The objective of this assignment is to build a data-driven inventory ordering policy "
    "for a single perishable SKU (art_id: 2921141) in a retail store. The policy must decide "
    "how many units to order each period to minimize the total accumulated cost over a 26-period "
    "live simulation (July 2 - August 2, 2021, excluding Sundays and official holidays)."
)

doc.add_heading("1.1 Cost Structure", level=2)

table = doc.add_table(rows=5, cols=3, style="Light Grid Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Parameter", "Value", "Interpretation"]
data = [
    ["Holding cost", "1 per unit/period", "Low cost for keeping inventory"],
    ["Shortage cost", "19 per unit", "Very high penalty for lost sales"],
    ["Expiry cost", "9 per unit", "Moderate penalty for waste"],
    ["Lead time", "2 periods", "Orders arrive 2 days after placement"],
]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r + 1].cells[c].text = val

doc.add_paragraph("")

doc.add_heading("1.2 Perishability and FIFO", level=2)
doc.add_paragraph(
    "Products have a shelf life of 2 periods. Units are tracked by age cohort: "
    "age-0 (fresh, delivered this period) and age-1 (delivered last period, expires at end of current period). "
    "When demand occurs, the oldest units (age-1) are sold first following a FIFO (First-In, First-Out) policy. "
    "Any age-1 units remaining unsold at the end of the period are discarded and incur the expiry cost."
)

doc.add_heading("1.3 Simulation Sequence", level=2)
doc.add_paragraph(
    "Each period follows this exact sequence:"
)
items = [
    "Start of period: Pipeline shipment arrives (becomes age-0 on-hand). "
    "Observe inventory state. Decide order quantity (enters pipeline, arrives in 2 periods).",
    "During the period: Customer demand occurs. Oldest units (age-1) are sold first (FIFO). "
    "Any demand that cannot be met incurs the shortage cost of 19 per unit.",
    "End of period: Remaining age-1 units expire (cost 9 each). "
    "A holding cost of 1 per unit is charged on all units still in inventory (before expiry removal). "
    "Age transition: age-0 units become age-1 for the next period.",
    "Move to next period: New shipments arrive.",
]
for i, item in enumerate(items, 1):
    doc.add_paragraph(f"{i}. {item}")

doc.add_heading("1.4 Starting Position", level=2)
doc.add_paragraph(
    "The game begins on 2021-07-02 with the starting inventory position [5, 4, 3]: "
    "5 units arriving the next day (in the pipeline), 4 units of age-0 (delivered today), "
    "and 3 units of age-1 (delivered yesterday, expiring at end of today). "
    "This gives 7 units on-hand and 5 in the pipeline."
)

doc.add_heading("1.5 Critical Ratio Analysis", level=2)
doc.add_paragraph(
    "The newsvendor critical ratio provides the theoretical foundation for our ordering policy. "
    "With an underage cost (shortage) of 19 and an overage cost (holding) of 1, the critical ratio is:"
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CR = Cu / (Cu + Co) = 19 / (19 + 1) = 0.95")
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph(
    "This means we should target the 95th percentile of the demand distribution: "
    "it is optimal to stock enough to satisfy demand 95% of the time. "
    "However, the perishability constraint (shelf life = 2) and expiry cost (9 per unit) "
    "mean that blindly ordering to the 95th percentile would cause excessive waste. "
    "Our policy must balance this tension."
)

doc.add_page_break()

# ============================================================
# 2. DATA EXPLORATION
# ============================================================
doc.add_heading("2. Data Exploration", level=1)

doc.add_heading("2.1 Dataset Overview", level=2)
doc.add_paragraph(
    "The training data (df_6_art_train_project.parquet) contains 7,143 rows covering "
    "6 SKUs from one retail store, spanning September 2017 to August 2021. "
    "Each row represents one selling day (Sundays and official holidays are excluded). "
    "The target SKU (art_id: 2921141, subgroup 109) has 1,191 rows, of which "
    "1,165 have known sales values (the remaining 26 are the test period with NaN sales)."
)

doc.add_heading("2.2 Features", level=2)
table = doc.add_table(rows=8, cols=2, style="Light Grid Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
feat_data = [
    ["Feature", "Description"],
    ["PROMO_01", "Binary indicator: 1 if product is on promotion, 0 otherwise"],
    ["PROMO_DEPTH", "Promotion discount depth (0 = no promo, up to 60%)"],
    ["PRC_2_norm", "Normalized price of the product"],
    ["OFFICIAL_HOLIDAY_01_f1", "Binary: 1 if tomorrow is an official holiday"],
    ["OFFICIAL_HOLIDAY_01_l1", "Binary: 1 if yesterday was an official holiday"],
    ["date", "Date of the observation (used to derive day-of-week)"],
    ["sales", "Target variable: number of units sold"],
]
for r, row_data in enumerate(feat_data):
    for c, val in enumerate(row_data):
        cell = table.rows[r].cells[c]
        cell.text = val
        if r == 0:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph("")

doc.add_heading("2.3 Target SKU Sales Statistics", level=2)
table = doc.add_table(rows=2, cols=6, style="Light Grid Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
stats_headers = ["Mean", "Median", "Std Dev", "Min", "Max", "Count"]
stats_values = ["5.00", "4.0", "4.28", "0", "24", "1,165"]
for c, h in enumerate(stats_headers):
    table.rows[0].cells[c].text = h
    table.rows[0].cells[c].paragraphs[0].runs[0].bold = True
for c, v in enumerate(stats_values):
    table.rows[1].cells[c].text = v

doc.add_paragraph("")

doc.add_heading("2.4 Key Demand Patterns", level=2)

doc.add_paragraph("Promotion effect:", style="List Bullet")
p = doc.add_paragraph(
    "Sales during promotions average 5.77 units/day vs. 4.33 without promotion. "
    "The test period features a heavy promotion (depth=40%) for the first 12 days, "
    "followed by 14 days without promotion. This means we expect higher demand early "
    "in the game and lower demand in the second half."
)

doc.add_paragraph("Day-of-week effect:", style="List Bullet")
p = doc.add_paragraph(
    "Saturday consistently shows the highest demand (~2x weekday average), "
    "followed by Friday. Monday through Thursday show lower, more stable demand. "
    "This weekly pattern is critical for ordering decisions since the 2-period lead time "
    "means Thursday orders serve Saturday demand."
)

doc.add_paragraph("Price sensitivity:", style="List Bullet")
p = doc.add_paragraph(
    "Lower prices (PRC_2_norm < 2.5) correlate with higher demand (mean ~6.3), "
    "independent of the promotion flag. The test period includes a stretch of very low "
    "price (1.99) at the end of July."
)

doc.add_heading("2.5 Test Period Structure (26 Periods)", level=2)

table = doc.add_table(rows=3, cols=4, style="Light Grid Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
tp_data = [
    ["Phase", "Dates", "Promo", "Expected Demand"],
    ["Phase 1 (12 days)", "Jul 2 - Jul 13", "Yes (depth=40%)", "Higher (~6-8/day)"],
    ["Phase 2 (14 days)", "Jul 14 - Aug 2", "No", "Normal (~3-5/day)"],
]
for r, row_data in enumerate(tp_data):
    for c, val in enumerate(row_data):
        cell = table.rows[r].cells[c]
        cell.text = val
        if r == 0:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph("")
doc.add_paragraph(
    "One holiday-adjacent period occurs around July 20-22. "
    "The transition from promo to non-promo around July 14 is a critical moment "
    "where demand drops suddenly and excess inventory may expire."
)

doc.add_page_break()

# ============================================================
# 3. DEMAND FORECASTING MODEL
# ============================================================
doc.add_heading("3. Demand Forecasting Model", level=1)

doc.add_heading("3.1 Approach: LightGBM Quantile Regression", level=2)
doc.add_paragraph(
    "Rather than predicting a single point forecast (which would not capture uncertainty), "
    "we train multiple LightGBM models, each targeting a different quantile of the "
    "conditional demand distribution. This directly answers the question: "
    "\"Given today's features (day-of-week, promo status, price, holidays), "
    "what is the Xth percentile of demand?\""
)

doc.add_paragraph(
    "We train 5 models in total:"
)
models_list = [
    "Poisson mean model: provides the expected (average) demand for reference.",
    "Quantile 0.50 (median): the central forecast, used as a floor for ordering.",
    "Quantile 0.75: upper-moderate demand scenario.",
    "Quantile 0.90: high demand scenario.",
    "Quantile 0.95: the primary ordering target, matching the critical ratio of 0.95.",
]
for m in models_list:
    doc.add_paragraph(m, style="List Bullet")

doc.add_heading("3.2 Feature Engineering", level=2)
doc.add_paragraph(
    "We engineer 8 features from the raw data:"
)
features_list = [
    "dow (0-6): Day of week extracted from the date. Captures weekly seasonality.",
    "is_saturday (binary): Explicit indicator for the highest-demand day.",
    "is_friday (binary): Explicit indicator for the second-highest-demand day.",
    "PROMO_01 (binary): Whether a promotion is active.",
    "PROMO_DEPTH (0-60): The discount percentage of the promotion.",
    "PRC_2_norm (continuous): Normalized product price.",
    "OFFICIAL_HOLIDAY_01_f1 (binary): Holiday tomorrow indicator.",
    "OFFICIAL_HOLIDAY_01_l1 (binary): Holiday yesterday indicator.",
]
for f in features_list:
    doc.add_paragraph(f, style="List Bullet")

doc.add_heading("3.3 Model Configuration", level=2)
doc.add_paragraph(
    "All LightGBM models share the same hyperparameters to ensure consistency:"
)
table = doc.add_table(rows=5, cols=2, style="Light Grid Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hp_data = [
    ["Hyperparameter", "Value"],
    ["num_leaves", "16 (prevents overfitting on 1,165 rows)"],
    ["min_child_samples", "20 (ensures stable leaf estimates)"],
    ["learning_rate", "0.05"],
    ["num_boost_round", "300"],
]
for r, row_data in enumerate(hp_data):
    for c, val in enumerate(row_data):
        cell = table.rows[r].cells[c]
        cell.text = val
        if r == 0:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph("")

doc.add_heading("3.4 Validation Results", level=2)
doc.add_paragraph(
    "We perform a time-series validation split using the last 200 training days as the test set "
    "(train: 965 rows, test: 200 rows). Results are evaluated using pinball loss and empirical coverage:"
)

table = doc.add_table(rows=5, cols=3, style="Light Grid Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
val_data = [
    ["Quantile", "Pinball Loss", "Actual Coverage (Target)"],
    ["0.50", "1.552", "64.0% (50%)"],
    ["0.75", "1.340", "79.5% (75%)"],
    ["0.90", "0.803", "90.0% (90%)"],
    ["0.95", "0.504", "93.0% (95%)"],
]
for r, row_data in enumerate(val_data):
    for c, val in enumerate(row_data):
        cell = table.rows[r].cells[c]
        cell.text = val
        if r == 0:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph("")
doc.add_paragraph(
    "The q90 model achieves exactly 90% coverage. The q95 model achieves 93%, slightly below "
    "the target but still well-calibrated. The q50 model over-covers at 64%, which is typical "
    "for count data with a long right tail - it means the median forecast tends to be slightly "
    "conservative, which is acceptable given our preference to avoid shortages."
)

doc.add_heading("3.5 Sample Forecasts for the Test Period", level=2)
doc.add_paragraph("Below are forecasts for the first 5 test days:")

table = doc.add_table(rows=6, cols=7, style="Light Grid Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
sf_data = [
    ["Date", "Day", "Promo", "Mean", "Q50", "Q90", "Q95"],
    ["Jul 02", "Fri", "Yes", "8.3", "8.2", "15.0", "16.4"],
    ["Jul 03", "Sat", "Yes", "7.9", "7.7", "13.6", "14.9"],
    ["Jul 05", "Mon", "Yes", "4.7", "5.2", "9.7", "9.9"],
    ["Jul 06", "Tue", "Yes", "4.2", "5.6", "6.6", "7.9"],
    ["Jul 07", "Wed", "Yes", "4.5", "5.1", "9.7", "11.8"],
]
for r, row_data in enumerate(sf_data):
    for c, val in enumerate(row_data):
        cell = table.rows[r].cells[c]
        cell.text = val
        if r == 0:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph("")
doc.add_paragraph(
    "Notice how Friday and Saturday show higher forecasts even within the promo period, "
    "confirming the importance of the day-of-week feature."
)

doc.add_page_break()

# ============================================================
# 4. INVENTORY SIMULATION ENGINE
# ============================================================
doc.add_heading("4. Inventory Simulation Engine", level=1)

doc.add_heading("4.1 State Representation", level=2)
doc.add_paragraph(
    "The inventory state is tracked using two arrays:"
)
doc.add_paragraph(
    "on_hand = [age_0, age_1]: Units currently in the store, tracked by freshness. "
    "age_0 units were delivered this period (fresh). age_1 units were delivered last period "
    "and will expire at the end of the current period.", style="List Bullet"
)
doc.add_paragraph(
    "pipeline = [arriving_next_period, arriving_in_2_periods]: Orders that have been placed "
    "but not yet delivered. pipeline[0] arrives tomorrow, pipeline[1] arrives the day after.",
    style="List Bullet"
)

doc.add_heading("4.2 Period Execution Logic", level=2)
doc.add_paragraph(
    "The engine implements the exact simulation sequence specified in the assignment. "
    "For each period, the step() function:"
)
steps = [
    "Receives the arriving shipment: pipeline[0] is added to on_hand as age_0.",
    "Shifts the pipeline: pipeline[0] = pipeline[1], pipeline[1] = new order.",
    "Processes demand via FIFO: sells age_1 units first, then age_0. Tracks shortages.",
    "Computes expiry: any remaining age_1 units are discarded (cost 9 each).",
    "Computes holding cost: all remaining units (age_0 + age_1) incur cost 1 each.",
    "Performs age transition: age_0 becomes age_1 for the next period. age_0 resets to 0.",
    "Returns a detailed result dict with all period metrics.",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("4.3 Cost Tracking", level=2)
doc.add_paragraph(
    "The engine maintains a running total cost and a full history log. "
    "Each period records: units arrived, ordered, demanded, sold (by age cohort), "
    "shortage, expired, holding units, all three cost components, and the resulting state. "
    "This provides complete visibility for debugging and analysis."
)

doc.add_page_break()

# ============================================================
# 5. ORDERING POLICY
# ============================================================
doc.add_heading("5. Ordering Policy", level=1)

doc.add_heading("5.1 Core Concept: Perishable Newsvendor with Forward Projection", level=2)
doc.add_paragraph(
    "Our ordering policy extends the classical newsvendor framework to account for "
    "perishability and lead time. The key insight is:"
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "An order placed at period t arrives at period t+2.\n"
    "Those units are age_0 at t+2 and age_1 at t+3 (then expire).\n"
    "So each order covers demand at t+2 and t+3."
)
run.italic = True

doc.add_paragraph("")
doc.add_paragraph(
    "The policy works in three steps: (1) forecast demand, (2) project forward inventory, "
    "(3) compute the order that closes the gap to the target quantile."
)

doc.add_heading("5.2 Step 1: Demand Forecasting", level=2)
doc.add_paragraph(
    "For each period, we generate quantile forecasts for the current period (t), "
    "the next period (t+1), and the arrival period (t+2) and beyond (t+3). "
    "These forecasts use the known features (day-of-week, promo status, price, holiday indicators) "
    "from the test dataset."
)

doc.add_heading("5.3 Step 2: Forward Inventory Projection", level=2)
doc.add_paragraph(
    "This is the most technically involved part of the policy. Before deciding how much to order, "
    "we need to estimate how much inventory will still be alive when the order arrives at t+2. "
    "We simulate forward through the FIFO and expiry mechanics:"
)

doc.add_paragraph(
    "Period t simulation: The current pipeline[0] arrives as age_0. "
    "Mean demand at t is consumed via FIFO (age_1 first, then age_0). "
    "Remaining age_1 expires. Surviving age_0 units carry forward as age_1.",
    style="List Bullet"
)
doc.add_paragraph(
    "Period t+1 simulation: pipeline[1] arrives as age_0. "
    "The carry-forward from t (now age_1) is consumed first by mean demand at t+1. "
    "Remaining age_1 expires. Surviving age_0 carries forward.",
    style="List Bullet"
)
doc.add_paragraph(
    "At t+2: The carry-forward from t+1 becomes age_1. "
    "The new order arrives as age_0. Together they must cover demand at t+2.",
    style="List Bullet"
)

doc.add_paragraph(
    "By simulating these mechanics, we get a precise estimate of 'existing_at_t2' - "
    "the inventory that will be alive at t+2 without our new order."
)

doc.add_heading("5.4 Step 3: Order Quantity Calculation", level=2)
doc.add_paragraph(
    "The order quantity is computed as:"
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("order = max(0, Q_target(t+2) - existing_at_t2)")
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph("")
doc.add_paragraph("Where Q_target is the target quantile of demand at t+2. Three adjustments are applied:")

doc.add_paragraph(
    "Target quantile selection (0.95 by default): The critical ratio of 19/20 = 0.95 "
    "means we target the 95th percentile of demand. This is intentionally aggressive "
    "because the cost of a shortage (19) far exceeds holding (1).",
    style="List Bullet"
)
doc.add_paragraph(
    "Maximum useful cap: The order cannot serve more than demand at t+2 and t+3 combined "
    "(since it expires after t+3). We cap: order <= Q_target(t+2) + Q_median(t+3) - existing.",
    style="List Bullet"
)
doc.add_paragraph(
    "Minimum floor: We always order at least enough to cover the median demand at t+2, "
    "providing a safety net even when projected inventory seems sufficient.",
    style="List Bullet"
)

doc.add_heading("5.5 End-Game Tapering", level=2)
doc.add_paragraph(
    "As the game approaches its final periods, orders that arrive late have fewer periods "
    "to be sold before the game ends. Any unsold units at the end represent pure waste. "
    "We progressively reduce the target quantile:"
)

table = doc.add_table(rows=5, cols=2, style="Light Grid Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
taper_data = [
    ["Periods Remaining After Arrival", "Target Quantile"],
    ["> 3 periods", "0.95 (full aggression)"],
    ["3 periods", "0.90"],
    ["2 periods", "0.75"],
    ["1 period or less", "0.50 (conservative)"],
]
for r, row_data in enumerate(taper_data):
    for c, val in enumerate(row_data):
        cell = table.rows[r].cells[c]
        cell.text = val
        if r == 0:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph("")
doc.add_paragraph(
    "This prevents the common pitfall of large end-game expiry costs, which in our backtests "
    "reduced total expiry costs by approximately 20%."
)

doc.add_page_break()

# ============================================================
# 6. BACKTEST RESULTS
# ============================================================
doc.add_heading("6. Backtest Results", level=1)

doc.add_paragraph(
    "We evaluated the policy under three deterministic demand scenarios to stress-test "
    "its behavior across different demand levels. Note: these are not realistic simulations "
    "(real demand varies period-to-period), but they reveal the policy's tendencies."
)

table = doc.add_table(rows=4, cols=5, style="Light Grid Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
bt_data = [
    ["Scenario", "Total Cost", "Holding", "Shortage", "Expiry"],
    ["Median demand", "375", "119", "76", "180"],
    ["Q75 demand", "639", "12", "627", "0"],
    ["Q90 demand", "1,522", "2", "1,520", "0"],
]
for r, row_data in enumerate(bt_data):
    for c, val in enumerate(row_data):
        cell = table.rows[r].cells[c]
        cell.text = val
        if r == 0:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph("")

doc.add_heading("6.1 Interpretation", level=2)

doc.add_paragraph(
    "Median scenario (375 total): This represents the most likely outcome. "
    "The cost is split between moderate holding (119), a small unavoidable shortage at period 2 (76), "
    "and expiry (180) mostly from the promo-to-non-promo transition and end-game.",
    style="List Bullet"
)
doc.add_paragraph(
    "Q75 scenario (639 total): When demand consistently exceeds the median, shortages dominate. "
    "This is expected - our policy targets Q95 of individual periods, not sustained high demand.",
    style="List Bullet"
)
doc.add_paragraph(
    "Q90 scenario (1,522 total): An extreme stress test. Consistent 90th-percentile demand "
    "will defeat any reasonable policy. The zero expiry confirms the policy doesn't over-order.",
    style="List Bullet"
)

doc.add_heading("6.2 Key Observations", level=2)
doc.add_paragraph(
    "Period 2 shortage is unavoidable: The starting pipeline has a gap (pipeline = [5, 0]), "
    "meaning no units arrive on day 2. With Saturday promo demand of ~8, a shortage of ~4 units "
    "is structural. Our policy orders 10 units on day 1, but they arrive on day 3.",
    style="List Bullet"
)
doc.add_paragraph(
    "Promo-to-non-promo transition: Around period 11 (July 14), the promotion ends and demand "
    "drops sharply. Inventory stocked for promo demand levels may expire. The policy handles this "
    "by using feature-aware forecasts that automatically reduce when promo=0.",
    style="List Bullet"
)
doc.add_paragraph(
    "Live game advantage: In the actual game, demand is revealed each period. The policy "
    "automatically adapts through the inventory state - high revealed demand depletes inventory, "
    "causing higher recommended orders next period. This reactive behavior is not captured in "
    "deterministic backtests.",
    style="List Bullet"
)

doc.add_page_break()

# ============================================================
# 7. HOW TO RUN
# ============================================================
doc.add_heading("7. How to Run the Live Game", level=1)

doc.add_heading("7.1 Prerequisites", level=2)
doc.add_paragraph("Python 3.8+ with the following packages: pandas, numpy, lightgbm, python-docx (for this report).")

doc.add_heading("7.2 Running the Game", level=2)
p = doc.add_paragraph()
run = p.add_run("python play_game.py")
run.font.name = "Consolas"
run.font.size = Pt(11)
run.bold = True

doc.add_paragraph("")
doc.add_paragraph(
    "The script loads in ~2 seconds, then presents an interactive prompt for each of the 26 periods. "
    "For each period, it displays:"
)
display_items = [
    "Current date, day-of-week, promo status, and price.",
    "On-hand inventory by age cohort: [age_0, age_1].",
    "Pipeline: [arriving next period, arriving in 2 periods].",
    "Demand forecast: mean, q50, q75, q90, q95.",
    "Recommended order quantity (computed by the policy).",
]
for item in display_items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph("")
doc.add_paragraph(
    "You can accept the recommended order (press Enter) or type a different quantity. "
    "After entering the revealed demand, the engine simulates the period and shows results "
    "including units sold, shortages, expiry, and the period cost breakdown."
)

doc.add_page_break()

# ============================================================
# 8. KEY DESIGN DECISIONS
# ============================================================
doc.add_heading("8. Key Design Decisions", level=1)

doc.add_heading("8.1 Why LightGBM Quantile Regression?", level=2)
doc.add_paragraph(
    "Several alternatives were considered:"
)
doc.add_paragraph(
    "Parametric distributions (Poisson, Negative Binomial): The data is overdispersed "
    "(variance/mean = 3.66), and demand has complex interactions (day-of-week x promo x price) "
    "that parametric models would miss without manual interaction terms.",
    style="List Bullet"
)
doc.add_paragraph(
    "Deep learning (LSTMs, transformers): With only 1,165 training rows and 8 features, "
    "deep models would overfit. The weak autocorrelation (0.17) means sequential models "
    "add little value over feature-based approaches.",
    style="List Bullet"
)
doc.add_paragraph(
    "LightGBM quantile regression wins because it: (a) directly outputs quantiles without "
    "distributional assumptions, (b) handles feature interactions automatically, "
    "(c) trains in under 1 second, and (d) is robust with small datasets when properly regularized.",
    style="List Bullet"
)

doc.add_heading("8.2 Why Not Reinforcement Learning?", level=2)
doc.add_paragraph(
    "With only 26 game periods and a 2-period look-ahead, the myopic newsvendor heuristic "
    "(adjusted for perishability) is near-optimal. The state space is small enough for an "
    "analytical policy. RL would require thousands of training episodes to converge and "
    "would be slower to execute during live play. The speed requirement (\"fast to compute, "
    "used live in class\") further favors our lightweight approach."
)

doc.add_heading("8.3 Why Single-SKU Training?", level=2)
doc.add_paragraph(
    "The dataset includes 5 additional SKUs. However, the two same-subgroup SKUs (109) have "
    "substantially different mean demand levels (6.86 and 2.89 vs. 5.00 for our target). "
    "Pooling would blur the demand signal without clear benefit, as 1,165 training rows is "
    "already sufficient for 8 features. Including different subgroup SKUs (125) would add "
    "noise from different product characteristics."
)

doc.add_heading("8.4 Cost Asymmetry Rationale", level=2)
doc.add_paragraph(
    "The entire policy is driven by the massive cost asymmetry: shortage (19) >> holding (1). "
    "This 19:1 ratio means that one unit of shortage costs the same as holding 19 units "
    "for one period. Even the expiry cost (9) is less than half the shortage cost. "
    "Therefore, our policy deliberately errs on the side of over-ordering. The only constraint "
    "on this aggression is the shelf life: units that cannot possibly be sold in 2 periods "
    "will expire at cost 9, which is still expensive enough to warrant the maximum-useful cap "
    "in our ordering logic."
)

doc.add_heading("8.5 Architecture Summary", level=2)
doc.add_paragraph(
    "The solution is organized into three clean, modular Python files:"
)
table = doc.add_table(rows=4, cols=3, style="Light Grid Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
arch_data = [
    ["File", "Purpose", "Key Class/Function"],
    ["demand_model.py", "Demand forecasting", "DemandForecaster (fit, predict_quantiles)"],
    ["inventory_engine.py", "Inventory simulation", "PerishableInventory (step, get_state)"],
    ["play_game.py", "Live game interface", "compute_recommended_order, main loop"],
]
for r, row_data in enumerate(arch_data):
    for c, val in enumerate(row_data):
        cell = table.rows[r].cells[c]
        cell.text = val
        if r == 0:
            cell.paragraphs[0].runs[0].bold = True

# ============================================================
# SAVE
# ============================================================
output_path = os.path.join(os.path.dirname(__file__), "Methodology_Report.docx")
doc.save(output_path)
print(f"Report saved to: {output_path}")
